from tkinter import Tk, Label, Button, Entry, messagebox, ttk, Canvas, PhotoImage, Listbox, simpledialog
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
import pandas as pd
import algoritmo
from algoritmo_archivos.base_datos import listar_ventas_bd, listar_ventas_por_cliente_bd, listar_ventas_por_producto_bd, anular_venta_bd, obtener_producto_por_codigo, actualizar_existencia_producto, listar_clientes_bd

buscar_cliente_label = None
buscar_cliente_entry = None
buscar_producto_label = None
buscar_producto_entry = None
lista_desplegable = None
tree = None
ventana_reporte_instance = None

#Crear carpetas si no existen
if not os.path.exists('word'):
    os.makedirs('word')
if not os.path.exists('excel'):
    os.makedirs('excel')

def enviar_correo(archivo, tipo):
    #Pedir correo electronico al usuario
    correo = simpledialog.askstring("Correo Electrónico", "Ingrese su correo electrónico:")

    if correo:
        #Configuración del correo
        remitente = "dignaciot@miumg.edu.gt"
        destinatario = correo
        asunto = f"Reporte en {tipo}"
        cuerpo = f"Reporte en {tipo}."

        #Crear el mensaje de correo
        mensaje = MIMEMultipart()
        mensaje['From'] = remitente
        mensaje['To'] = destinatario
        mensaje['Subject'] = asunto

        mensaje.attach(MIMEText(cuerpo, 'plain'))

        #Adjuntar el archivo
        adjunto = open(archivo, "rb")
        parte = MIMEBase('application', 'octet-stream')
        parte.set_payload((adjunto).read())
        encoders.encode_base64(parte)
        parte.add_header('Content-Disposition', f"attachment; filename= {os.path.basename(archivo)}")
        mensaje.attach(parte)

        #Enviar el correo
        servidor = smtplib.SMTP('smtp.gmail.com', 587)
        servidor.starttls()
        servidor.login(remitente.encode('utf-8'), "tu_contraseña".encode('utf-8'))
        texto = mensaje.as_string()
        servidor.sendmail(remitente, destinatario, texto)
        servidor.quit()

def generar_reporte_word():
    #Crear documento en Word
    doc = Document()
    doc.add_heading('Reporte', 0)

    #Agregar datos de la tabla al documento de Word
    table = doc.add_table(rows=1, cols=len(tree["columns"]))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(tree["columns"]):
        hdr_cells[i].text = col

    for row in tree.get_children():
        row_data = tree.item(row)['values']
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

    #Guardar el archivo en la carpeta 'word'
    archivo_word = os.path.join('word', 'reporte.docx')
    doc.save(archivo_word)

    #Enviar el archivo por correo electronico
    enviar_correo(archivo_word, "Word")

def generar_reporte_excel():
    #Obtener datos de la tabla y convertirlos en un DataFrame de pandas
    data = []
    for row in tree.get_children():
        row_data = tree.item(row)['values']
        data.append(row_data)
    
    df = pd.DataFrame(data, columns=tree["columns"])

    #Guardar el archivo en la carpeta 'excel'
    archivo_excel = os.path.join('excel', 'reporte.xlsx')
    df.to_excel(archivo_excel, index=False)

    #Enviar el archivo por correo electrónico
    enviar_correo(archivo_excel, "Excel")

def ocultar_entry_productos():
    global buscar_producto_label, buscar_producto_entry, lista_desplegable
    if buscar_producto_label:
        buscar_producto_label.destroy()
    if buscar_producto_entry:
        buscar_producto_entry.destroy()
    if lista_desplegable:
        lista_desplegable.place_forget()

def ocultar_entry_clientes():
    global buscar_cliente_label, buscar_cliente_entry, lista_desplegable
    if buscar_cliente_label:
        buscar_cliente_label.destroy()
    if buscar_cliente_entry:
        buscar_cliente_entry.destroy()
    if lista_desplegable:
        lista_desplegable.place_forget()

def ejecutar_ambos_clientes(tree, ventana_reporte_instance, boton_anular_compra):
    ocultar_entry_productos()
    mostrar_ventas_cliente(tree, ventana_reporte_instance, boton_anular_compra)

def ejecutar_ambos_productos(tree, ventana_reporte_instance, boton_anular_compra):
    ocultar_entry_clientes()
    mostrar_ventas_producto(tree, ventana_reporte_instance, boton_anular_compra)

#Funcion para mostrar ventas generales
def mostrar_ventas_generales(tree, boton_anular_compra):
    ocultar_entry_clientes()
    ocultar_entry_productos()

    boton_anular_compra.place(x=450, y=620)
    for item in tree.get_children():
        tree.delete(item)
    
    tree["columns"] = ("fecha", "producto", "codigo_producto", "cantidad", "cliente", "codigo_cliente", "total")
    tree.heading("fecha", text="Fecha")
    tree.heading("producto", text="Producto")
    tree.heading("codigo_producto", text="Código")
    tree.heading("cantidad", text="Cantidad Vendida")
    tree.heading("cliente", text="Cliente")
    tree.heading("codigo_cliente", text="Código")
    tree.heading("total", text="Total")
    tree.column("fecha", anchor="center", width=150)
    tree.column("producto", anchor="center", width=150)
    tree.column("codigo_producto", anchor="center", width=50)
    tree.column("cantidad", anchor="center", width=80)
    tree.column("cliente", anchor="center", width=100)
    tree.column("codigo_cliente", anchor="center", width=50)
    tree.column("total", anchor="center", width=100)

    for row in listar_ventas_bd():
        tree.insert("", "end", values=row)

#Funcion para mostrar ventas por cliente
def mostrar_ventas_cliente(tree, ventana_reporte_instance, boton_anular_compra):
    global buscar_cliente_label
    global buscar_cliente_entry
    global lista_desplegable

    boton_anular_compra.place_forget()
    for item in tree.get_children():
        tree.delete(item)
    
    tree["columns"] = ("fecha", "codigo_cliente", "cliente", "total")
    tree.heading("fecha", text="Fecha")
    tree.heading("codigo_cliente", text="ID Cliente")
    tree.heading("cliente", text="Cliente")
    tree.heading("total", text="Total")
    tree.column("fecha", anchor="center", width=150)
    tree.column("codigo_cliente", anchor="center", width=100)
    tree.column("cliente", anchor="center", width=200)
    tree.column("total", anchor="center", width=100)

    #Crear entrada para buscar cliente
    buscar_cliente_label = Label(ventana_reporte_instance, text="Buscar Cliente:", font=("Arial", 12), bg="SkyBlue3")
    buscar_cliente_label.place(x=20, y=350)
    buscar_cliente_entry = Entry(ventana_reporte_instance)
    buscar_cliente_entry.place(x=20, y=380)

    #Crear lista desplegable para autocompletado
    lista_desplegable = Listbox(ventana_reporte_instance)

    def buscar_cliente(event):
        cliente_input = buscar_cliente_entry.get()
        lista_desplegable.delete(0, 'end')
        for row in listar_clientes_bd():
            if cliente_input.lower() in row[1].lower():  
                lista_desplegable.insert('end', row[1])  
        lista_desplegable.place(x=20, y=410)

    def seleccionar_cliente(event):
        seleccionado = lista_desplegable.get(lista_desplegable.curselection())
        buscar_cliente_entry.delete(0, 'end')
        buscar_cliente_entry.insert(0, seleccionado)
        lista_desplegable.place_forget()
        for item in tree.get_children():
            tree.delete(item)
        for row in listar_ventas_por_cliente_bd(seleccionado):
            tree.insert("", "end", values=row)

    lista_desplegable.bind("<<ListboxSelect>>", seleccionar_cliente)
    buscar_cliente_entry.bind("<KeyRelease>", buscar_cliente)

#Funcion para mostrar ventas por producto
def mostrar_ventas_producto(tree, ventana_reporte_instance, boton_anular_compra):
    global buscar_producto_label
    global buscar_producto_entry
    global lista_desplegable

    boton_anular_compra.place_forget()
    for item in tree.get_children():
        tree.delete(item)
    
        tree["columns"] = ("fecha", "codigo_producto", "producto", "cantidad", "total")
    tree.heading("fecha", text="Fecha")
    tree.heading("codigo_producto", text="ID Producto")
    tree.heading("producto", text="Producto")
    tree.heading("cantidad", text="Cantidad")
    tree.heading("total", text="Total")
    tree.column("fecha", anchor="center", width=150)
    tree.column("codigo_producto", anchor="center", width=100)
    tree.column("producto", anchor="center", width=200)
    tree.column("cantidad", anchor="center", width=100)
    tree.column("total", anchor="center", width=100)

    #Crear entrada para buscar producto
    buscar_producto_label = Label(ventana_reporte_instance, text="Buscar Producto:", font=("Arial", 12), bg="SkyBlue3")
    buscar_producto_label.place(x=20, y=350)
    buscar_producto_entry = Entry(ventana_reporte_instance)
    buscar_producto_entry.place(x=20, y=380)

    #Crear lista desplegable para autocompletado
    lista_desplegable = Listbox(ventana_reporte_instance)

    def buscar_producto(event):
        producto_input = buscar_producto_entry.get()
        lista_desplegable.delete(0, 'end')
        for row in listar_ventas_por_producto_bd(producto_input):
            lista_desplegable.insert('end', row[2])  
        lista_desplegable.place(x=20, y=410)

    def seleccionar_producto(event):
        seleccionado = lista_desplegable.get(lista_desplegable.curselection())
        buscar_producto_entry.delete(0, 'end')
        buscar_producto_entry.insert(0, seleccionado)
        lista_desplegable.place_forget()
        for item in tree.get_children():
            tree.delete(item)
        for row in listar_ventas_por_producto_bd(seleccionado):
            tree.insert("", "end", values=row)

    lista_desplegable.bind("<<ListboxSelect>>", seleccionar_producto)
    buscar_producto_entry.bind("<KeyRelease>", buscar_producto)

def ventana_reporte():
    global tree
    global ventana_reporte_instance
    if ventana_reporte_instance is not None:
        ventana_reporte_instance.destroy()
    ventana_reporte_instance = Tk()
    ventana_reporte_instance.title("Reporte")
    ventana_reporte_instance.geometry("1100x700+100+0")
    ventana_reporte_instance.resizable(False, False)

    canvas_bg = Canvas(ventana_reporte_instance, width=1100, height=700)
    canvas_bg.pack(fill="both", expand=True)

    #Crear el degradado
    for i in range(700):
        ratio = i / 700
        color = '#%02x%02x%02x' % (
            int(45 + ratio * (149 - 45)),
            int(106 + ratio * (213 - 106)),
            int(79 + ratio * (178 - 79))
        )
        canvas_bg.create_rectangle(0, i, 1100, i + 1, fill=color, outline=color)

    #Frame para los botones de seleccion de reporte
    frame_botones_seleccion = ttk.Frame(ventana_reporte_instance, padding="10")
    frame_botones_seleccion.place(x=300, y=20)

    #Boton para mostrar ventas generales

    ruta_imagen5 = os.path.join("algoritmo_archivos", "iconos", "general.png")
    general_img = PhotoImage(file=ruta_imagen5)
    boton_ventas_generales = Button(frame_botones_seleccion, text=" Ventas Generales", command=lambda: mostrar_ventas_generales(tree, boton_anular_compra))
    boton_ventas_generales.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 10, "bold"), image=general_img, compound="left")
    boton_ventas_generales.grid(row=0, column=0, padx=10)

    #Boton para mostrar ventas por cliente

    ruta_imagen6 = os.path.join("algoritmo_archivos", "iconos", "cliente.png")
    cliente_img = PhotoImage(file=ruta_imagen6)
    boton_ventas_cliente = Button(frame_botones_seleccion, text=" Ventas por Cliente", command=lambda: ejecutar_ambos_clientes(tree, ventana_reporte_instance, boton_anular_compra))
    boton_ventas_cliente.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 10, "bold"), image=cliente_img, compound="left")
    boton_ventas_cliente.grid(row=0, column=1, padx=10)

    #Boton para mostrar ventas por producto

    ruta_imagen7 = os.path.join("algoritmo_archivos", "iconos", "producto.png")
    producto_img = PhotoImage(file=ruta_imagen7)
    boton_ventas_producto = Button(frame_botones_seleccion, text=" Ventas por Producto", command=lambda: ejecutar_ambos_productos(tree, ventana_reporte_instance, boton_anular_compra))
    boton_ventas_producto.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 10, "bold"), image=producto_img, compound="left")
    boton_ventas_producto.grid(row=0, column=2, padx=10)

    #Boton para generar reporte en Word
    ruta_imagen2 = os.path.join("algoritmo_archivos", "iconos", "word.png")
    word_img = PhotoImage(file=ruta_imagen2)
    boton_generar_word = Button(ventana_reporte_instance, text="  Generar Reporte en Word", command=generar_reporte_word)
    boton_generar_word.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 10, "bold"), image=word_img, compound="left")
    boton_generar_word.place(x=20, y=170)

    #Boton para generar reporte en Excel
    ruta_imagen3 = os.path.join("algoritmo_archivos", "iconos", "excel.png")
    excel_img = PhotoImage(file=ruta_imagen3)
    boton_generar_excel = Button(ventana_reporte_instance, text="  Generar Reporte en Excel", command=generar_reporte_excel)
    boton_generar_excel.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 10, "bold"), image=excel_img, compound="left")
    boton_generar_excel.place(x=20, y=230)

    #Tabla para mostrar los reportes
    columns = ("fecha", "producto", "codigo_producto", "cantidad", "cliente", "codigo_cliente", "total")
    tree = ttk.Treeview(ventana_reporte_instance, columns=columns, show="headings")
    tree.heading("fecha", text="Fecha")
    tree.heading("producto", text="Producto")
    tree.heading("codigo_producto", text="Código")
    tree.heading("cantidad", text="Cantidad Vendida")
    tree.heading("cliente", text="Cliente")
    tree.heading("codigo_cliente", text="Código")
    tree.heading("total", text="Total")
    tree.column("fecha", anchor="center", width=150)
    tree.column("producto", anchor="center", width=150)
    tree.column("codigo_producto", anchor="center", width=50)
    tree.column("cantidad", anchor="center", width=80)
    tree.column("cliente", anchor="center", width=100)
    tree.column("codigo_cliente", anchor="center", width=50)
    tree.column("total", anchor="center", width=100)
    tree.place(x=250, y=100, width=700, height=500)

    #Funcion para cargar datos desde la base de datos
    def cargar_datos():
        for item in tree.get_children():
            tree.delete(item)
        
        for row in listar_ventas_bd():
            tree.insert("", "end", values=row)

    cargar_datos()

    #Funcion para regresar a la ventana principal
    def regresar_algoritmo():
        ventana_reporte_instance.destroy()
        algoritmo.abrir_ventana()

    ruta_imagen4 = os.path.join("algoritmo_archivos", "iconos", "regresar_24.png")
    regresar_img = PhotoImage(file=ruta_imagen4)
    boton_regresar = Button(ventana_reporte_instance, text="  REGRESAR", command=regresar_algoritmo)
    boton_regresar.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 9, "bold"), image=regresar_img, compound="left")
    boton_regresar.place(x=50, y=620)

    #Boton para anular compra
    def anular_venta():
        selected_item = tree.selection()
        if selected_item:
            respuesta = messagebox.askyesno("Confirmación", "¿Estás seguro de que deseas anular esta venta?")
            if respuesta:
                item = tree.item(selected_item)
                codigo_producto = int(item['values'][2])
                cantidad_vendida = int(item['values'][3])
                codigo_cliente = item['values'][5]
                
                try:
                    #Anular la venta en la base de datos
                    anular_venta_bd(codigo_producto, codigo_cliente)
                    
                    #Actualizar el inventario
                    producto = obtener_producto_por_codigo(codigo_producto)
                    if producto:
                        nueva_existencia = producto[2] + cantidad_vendida
                        actualizar_existencia_producto(codigo_producto, nueva_existencia)
                    
                    # Eliminar el item de la lista
                    tree.delete(selected_item)
                    messagebox.showinfo("Éxito", "Venta anulada exitosamente")
                except Exception as e:
                    messagebox.showerror("Error", f"Error al anular la venta: {e}")
        else:
            messagebox.showwarning("Advertencia", "Por favor seleccione una venta para anular")

    ruta_imagen1 = os.path.join("algoritmo_archivos", "iconos", "eliminar.png")
    eliminar_img = PhotoImage(file=ruta_imagen1)
    boton_anular_compra = Button(ventana_reporte_instance, text=" Anular Venta", command=anular_venta)
    boton_anular_compra.config(bg="#FEFBEA", bd=10, relief="raised", font=("Helvetica", 14, "bold"), image=eliminar_img, compound="left")
    boton_anular_compra.place(x=520, y=620)

    ventana_reporte_instance.mainloop()

if __name__ == "__main__":
    ventana_reporte()
